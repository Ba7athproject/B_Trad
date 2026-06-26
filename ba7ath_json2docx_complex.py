import sys
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

def is_arabic(s):
    # Détecte de l'arabe unicode dans une chaîne
    return any('\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F' for c in s)

def add_arabic_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run._element.rPr.set(qn('w:rtl'), "1")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.style = style
    return p

def stylize_paragraph(par, style, align=None):
    par.style = style
    if align:
        par.alignment = align

def add_complex_blocks(doc, blocks):
    for blk in blocks:
        value = blk.get("corrige") or blk.get("trad") or blk.get("text", "")
        if blk["type"] == "Heading":
            # Heading traduit avec détection AR
            if is_arabic(value):
                add_arabic_paragraph(doc, value, style="Heading 1")
            else:
                doc.add_heading(value, level=1)
        elif blk["type"] == "Subtitle":
            if is_arabic(value):
                add_arabic_paragraph(doc, value, style="Subtitle")
            else:
                p = doc.add_paragraph(value)
                stylize_paragraph(p, 'Subtitle')
        elif blk["type"] == "Table":
            table = doc.add_table(rows=1, cols=len(blk["columns"]))
            table.style = "Table Grid"
            for i, col in enumerate(blk["columns"]): table.cell(0, i).text = str(col)
            for row in blk["data"]:
                row_cells = table.add_row().cells
                for i, val in enumerate(row): row_cells[i].text = str(val)
        elif blk["type"] == "List":
            for item in blk.get("items", [value]):
                if is_arabic(item):
                    add_arabic_paragraph(doc, item, style='List Bullet')
                else:
                    doc.add_paragraph(item, style='List Bullet')
        elif blk["type"] == "OCR":
            doc.add_paragraph("[Bloc OCR]", style='Intense Quote')
            # OCR arabic ou autre, détection directe
            if is_arabic(value):
                add_arabic_paragraph(doc, value, style='Quote')
            else:
                doc.add_paragraph(value, style='Quote')
        elif blk["type"] == "Paragraph":
            if is_arabic(value):
                add_arabic_paragraph(doc, value, style='Normal')
            else:
                par = doc.add_paragraph(value)
                stylize_paragraph(par, 'Normal')
        else:
            # Fallback
            if is_arabic(value):
                add_arabic_paragraph(doc, value, style='Normal')
            else:
                par = doc.add_paragraph(value)
                stylize_paragraph(par, 'Normal')

def json_to_docx_complex(json_in, docx_out):
    with open(json_in, encoding="utf-8") as f:
        pages = json.load(f)
    doc = Document()
    for page in pages:
        doc.add_heading(f"Page {page['page']}", level=1)
        add_complex_blocks(doc, page["blocks"])
        doc.add_page_break()
    doc.save(docx_out)
    print(f"[Mise en page complexe] {json_in} → {docx_out}")

if __name__ == "__main__":
    json_in, docx_out = sys.argv[1], sys.argv[2]
    json_to_docx_complex(json_in, docx_out)
